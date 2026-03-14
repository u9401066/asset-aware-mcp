"""
Tests for MarkdownDocxConverter and export_from_markdown service.

Covers:
- Markdown → DOCX conversion (headings, paragraphs, lists, tables, etc.)
- MarkdownDocxConverter inline formatting
- Application-layer export_from_markdown orchestration
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.infrastructure.markdown_converter import MarkdownDocxConverter


@pytest.fixture
def converter():
    return MarkdownDocxConverter()


# ============================================================================
# Basic conversion
# ============================================================================


class TestBasicConversion:
    def test_empty_markdown(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "empty.docx"
        result = converter.convert("", out)
        assert result.exists()
        assert result.stat().st_size > 0

    def test_single_paragraph(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "para.docx"
        converter.convert("Hello world", out)
        assert out.exists()

    def test_heading_levels(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "# H1\n## H2\n### H3\n#### H4\n##### H5\n###### H6"
        out = tmp_path / "headings.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_output_path_returned(
        self, converter: MarkdownDocxConverter, tmp_path: Path
    ):
        out = tmp_path / "test.docx"
        result = converter.convert("Test", out)
        assert result == out


# ============================================================================
# Inline formatting
# ============================================================================


class TestInlineFormatting:
    def test_bold(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "bold.docx"
        converter.convert("This is **bold** text", out)
        assert out.exists()

    def test_italic(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "italic.docx"
        converter.convert("This is *italic* text", out)
        assert out.exists()

    def test_bold_italic(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "bi.docx"
        converter.convert("This is ***bold italic*** text", out)
        assert out.exists()

    def test_strikethrough(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "strike.docx"
        converter.convert("This is ~~struck~~ text", out)
        assert out.exists()

    def test_inline_code(self, converter: MarkdownDocxConverter, tmp_path: Path):
        out = tmp_path / "code.docx"
        converter.convert("Use `print()` for output", out)
        assert out.exists()


# ============================================================================
# Block elements
# ============================================================================


class TestBlockElements:
    def test_code_block(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "```python\nprint('hello')\nx = 42\n```"
        out = tmp_path / "codeblock.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_unordered_list(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "- Item 1\n- Item 2\n- Item 3"
        out = tmp_path / "ul.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_ordered_list(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "1. First\n2. Second\n3. Third"
        out = tmp_path / "ol.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_blockquote(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "> This is a quote\n> with two lines"
        out = tmp_path / "quote.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_horizontal_rule(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "Above\n\n---\n\nBelow"
        out = tmp_path / "hr.docx"
        converter.convert(md, out)
        assert out.exists()


# ============================================================================
# Table support
# ============================================================================


class TestTables:
    def test_simple_table(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = "| Name | Age |\n| --- | --- |\n| Alice | 30 |\n| Bob | 25 |"
        out = tmp_path / "table.docx"
        converter.convert(md, out)
        assert out.exists()

    def test_table_with_br(self, converter: MarkdownDocxConverter, tmp_path: Path):
        """Tables with <br> should have newlines in cells."""
        md = "| Col |\n| --- |\n| Line 1<br>Line 2 |"
        out = tmp_path / "br_table.docx"
        converter.convert(md, out)
        assert out.exists()


# ============================================================================
# Complex document
# ============================================================================


class TestComplexDocument:
    def test_mixed_content(self, converter: MarkdownDocxConverter, tmp_path: Path):
        md = """# Report Title

## Introduction

This is a **bold** introduction with *italic* and `code`.

### Key Findings

1. First finding
2. Second finding
3. Third finding

| Metric | Value |
| --- | --- |
| Score | 95 |
| Grade | A |

> Important note: this is a blockquote.

---

## Conclusion

- Point A
- Point B
- Point C

```python
def main():
    print("Hello")
```
"""
        out = tmp_path / "complex.docx"
        converter.convert(md, out)
        assert out.exists()
        assert out.stat().st_size > 5000  # Non-trivial size

    def test_crlf_line_endings(self, converter: MarkdownDocxConverter, tmp_path: Path):
        """Windows CRLF line endings should be handled correctly."""
        md = "# Title\r\n\r\nParagraph one.\r\n\r\n## Section\r\n\r\n- item 1\r\n- item 2\r\n"
        out = tmp_path / "crlf.docx"
        converter.convert(md, out)
        doc = Document(str(out))
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "Title" in headings
        assert "Section" in headings

    def test_bold_not_greedy_across_stars(
        self, converter: MarkdownDocxConverter, tmp_path: Path
    ):
        """Ensure **bold** does not consume ***bold_italic*** markers."""
        md = "This is ***bold italic*** and **bold** text"
        out = tmp_path / "greedy.docx"
        converter.convert(md, out)
        doc = Document(str(out))
        runs = doc.paragraphs[0].runs
        has_bold_italic = any(r.bold and r.italic for r in runs)
        has_bold_only = any(r.bold and not r.italic for r in runs)
        assert has_bold_italic
        assert has_bold_only


# ============================================================================
# Application layer: export_from_markdown
# ============================================================================


class TestExportFromMarkdown:
    @pytest.fixture
    def mock_service(self, tmp_path):
        """Create a DocxService with a mock repository."""
        from unittest.mock import MagicMock

        from src.application.docx_service import DocxService

        repo = MagicMock()
        return DocxService(repo)

    @pytest.mark.asyncio
    async def test_export_docx_from_text(self, mock_service, tmp_path):
        out = tmp_path / "export.docx"
        result = await mock_service.export_from_markdown(
            md_text="# Hello\n\nWorld",
            output_path=str(out),
            output_format="docx",
        )
        assert result["success"] is True
        assert Path(result["output_path"]).exists()

    @pytest.mark.asyncio
    async def test_export_docx_from_file(self, mock_service, tmp_path):
        md_file = tmp_path / "input.md"
        md_file.write_text("# Test\n\nContent here", encoding="utf-8")
        out = tmp_path / "from_file.docx"
        result = await mock_service.export_from_markdown(
            md_path=str(md_file),
            output_path=str(out),
            output_format="docx",
        )
        assert result["success"] is True
        assert Path(result["output_path"]).exists()

    @pytest.mark.asyncio
    async def test_export_missing_file(self, mock_service):
        result = await mock_service.export_from_markdown(
            md_path="/nonexistent/file.md",
            output_format="docx",
        )
        assert result["success"] is False
        assert "not found" in result["error"]

    @pytest.mark.asyncio
    async def test_export_no_content(self, mock_service):
        result = await mock_service.export_from_markdown(output_format="docx")
        assert result["success"] is False
        assert "No markdown" in result["error"]

    @pytest.mark.asyncio
    async def test_export_invalid_format(self, mock_service):
        result = await mock_service.export_from_markdown(
            md_text="Hello", output_format="html"
        )
        assert result["success"] is False
        assert "Unsupported" in result["error"]

    @pytest.mark.asyncio
    async def test_export_default_output_path_from_md_path(
        self, mock_service, tmp_path
    ):
        md_file = tmp_path / "notes.md"
        md_file.write_text("# Notes", encoding="utf-8")
        result = await mock_service.export_from_markdown(
            md_path=str(md_file),
            output_format="docx",
        )
        assert result["success"] is True
        assert result["output_path"].endswith(".docx")
