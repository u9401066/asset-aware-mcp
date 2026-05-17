"""Unit tests for DocumentService CRUD and PDF→DOCX conversion helpers."""

from __future__ import annotations

import io
import json
from hashlib import sha256
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock

import pytest
from PIL import Image, ImageDraw

from src.application.document_service import (
    DocumentService,
    build_doc_id_unique_suffix,
    format_page_ranges,
    normalize_page_ranges,
    remap_markdown_page_markers,
)
from src.application.markdown_block_builder import build_markdown_blocks
from src.domain.entities import (
    DocumentAssets,
    DocumentManifest,
    IngestResult,
    TableAsset,
)
from src.domain.etl_profile import ETLProfile
from src.infrastructure.file_storage import FileStorage
from src.infrastructure.marker_adapter import MarkerBlock, MarkerParseResult


def _test_png_bytes(label: str, *, size: tuple[int, int] = (140, 110)) -> bytes:
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((8, 8, size[0] - 8, size[1] - 8), outline="black", width=4)
    draw.line((12, 12, size[0] - 12, size[1] - 12), fill="black", width=3)
    draw.text((18, 18), label, fill="black")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def test_normalize_page_ranges_merges_adjacent_ranges() -> None:
    assert normalize_page_ranges(["1-3", "4", "8-9", "9-10"], 12) == (
        (1, 4),
        (8, 10),
    )


def test_normalize_page_ranges_rejects_out_of_bounds_pages() -> None:
    with pytest.raises(ValueError, match="exceeds total page count"):
        normalize_page_ranges(["3-12"], 10)


def test_remap_markdown_page_markers_rewrites_subset_numbers() -> None:
    markdown = "<!-- Page 1 -->\nA\n<!-- Page 2 -->\nB"
    remapped = remap_markdown_page_markers(markdown, [10, 12])
    assert remapped == "<!-- Page 10 -->\nA\n<!-- Page 12 -->\nB"


def test_build_doc_id_unique_suffix_includes_page_ranges() -> None:
    suffix = build_doc_id_unique_suffix(Path("paper.pdf"), ((3, 5), (9, 9)))
    assert suffix.endswith("#pages=3-5,9")
    assert format_page_ranges(((3, 5), (9, 9))) == "3-5,9"


@pytest.mark.asyncio
async def test_delete_document_success() -> None:
    repository = MagicMock()
    repository.load_manifest.return_value = MagicMock(filename="paper.pdf")
    repository.delete_document.return_value = True

    service = DocumentService(repository=repository, pdf_extractor=MagicMock())

    result = await service.delete_document("doc_123")

    assert result == {
        "success": True,
        "doc_id": "doc_123",
        "filename": "paper.pdf",
        "warnings": [],
    }


@pytest.mark.asyncio
async def test_delete_document_removes_kg_entries() -> None:
    repository = MagicMock()
    repository.load_manifest.return_value = MagicMock(filename="paper.pdf")
    repository.delete_document.return_value = True
    knowledge_graph = MagicMock()
    knowledge_graph.is_available = True
    knowledge_graph.delete_document = AsyncMock(
        return_value={"status": "success", "message": "deleted"}
    )

    service = DocumentService(
        repository=repository,
        pdf_extractor=MagicMock(),
        knowledge_graph=knowledge_graph,
    )

    result = await service.delete_document("doc_123")

    assert result["success"] is True
    assert result["warnings"] == []
    assert result["knowledge_graph_status"] == "success"
    knowledge_graph.delete_document.assert_awaited_once_with("doc_123")


@pytest.mark.asyncio
async def test_delete_document_warns_when_kg_delete_fails() -> None:
    repository = MagicMock()
    repository.load_manifest.return_value = MagicMock(filename="paper.pdf")
    repository.delete_document.return_value = True
    knowledge_graph = MagicMock()
    knowledge_graph.is_available = True
    knowledge_graph.delete_document = AsyncMock(side_effect=RuntimeError("boom"))

    service = DocumentService(
        repository=repository,
        pdf_extractor=MagicMock(),
        knowledge_graph=knowledge_graph,
    )

    result = await service.delete_document("doc_123")

    assert result["success"] is True
    assert result["warnings"]
    assert "Knowledge graph deletion failed" in result["warnings"][0]


@pytest.mark.asyncio
async def test_ingest_can_skip_knowledge_graph_indexing(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")
    repository = FileStorage(tmp_path / "data")
    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1
    pdf_extractor.extract_text.return_value = "# Title\n\nBody"
    pdf_extractor.extract_images.return_value = []
    pdf_extractor.extract_tables.return_value = []
    pdf_extractor.get_toc.return_value = []
    pdf_extractor.get_title.return_value = ""
    knowledge_graph = MagicMock()
    knowledge_graph.is_available = True
    knowledge_graph.insert = AsyncMock(
        side_effect=AssertionError("KG insert should be skipped")
    )
    knowledge_graph.extract_entities = AsyncMock(
        side_effect=AssertionError("KG entity extraction should be skipped")
    )

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        knowledge_graph=knowledge_graph,
    )
    monkeypatch.setattr(service, "_extract_and_save_images", AsyncMock(return_value=[]))
    monkeypatch.setattr(service, "_extract_tables", AsyncMock(return_value=[]))

    results = await service.ingest(
        [str(pdf_path)],
        index_knowledge_graph=False,
    )

    assert results[0].success is True
    knowledge_graph.insert.assert_not_awaited()
    knowledge_graph.extract_entities.assert_not_awaited()


@pytest.mark.asyncio
async def test_ingest_skips_knowledge_graph_by_default(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")
    repository = FileStorage(tmp_path / "data")
    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1
    pdf_extractor.extract_text.return_value = "# Title\n\nBody"
    pdf_extractor.extract_images.return_value = []
    pdf_extractor.extract_tables.return_value = []
    pdf_extractor.get_toc.return_value = []
    pdf_extractor.get_title.return_value = ""
    knowledge_graph = MagicMock()
    knowledge_graph.is_available = True
    knowledge_graph.insert = AsyncMock(
        side_effect=AssertionError("KG insert should require an explicit opt-in")
    )
    knowledge_graph.extract_entities = AsyncMock(
        side_effect=AssertionError(
            "KG entity extraction should require an explicit opt-in"
        )
    )

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        knowledge_graph=knowledge_graph,
    )
    monkeypatch.setattr(service, "_extract_and_save_images", AsyncMock(return_value=[]))
    monkeypatch.setattr(service, "_extract_tables", AsyncMock(return_value=[]))

    results = await service.ingest([str(pdf_path)])

    assert results[0].success is True
    knowledge_graph.insert.assert_not_awaited()
    knowledge_graph.extract_entities.assert_not_awaited()


@pytest.mark.asyncio
async def test_convert_pdf_to_docx_rejects_fidelity_mode() -> None:
    service = DocumentService(repository=MagicMock(), pdf_extractor=MagicMock())

    result = await service.convert_pdf_to_docx("doc_123", mode="fidelity")

    assert result["success"] is False
    assert "content mode only" in result["error"]


@pytest.mark.asyncio
async def test_convert_pdf_to_docx_success(monkeypatch, tmp_path: Path) -> None:
    repository = MagicMock()
    repository.get_doc_dir.return_value = tmp_path
    repository.load_markdown.return_value = "# Title\n\nHello world"
    repository.load_manifest.return_value = MagicMock(
        title="Paper Title",
        assets=MagicMock(figures=[], tables=[MagicMock()]),
    )

    service = DocumentService(repository=repository, pdf_extractor=MagicMock())

    captured: dict[str, object] = {}

    def fake_build(markdown, manifest, output_path):
        captured["markdown"] = markdown
        captured["output_path"] = output_path

    monkeypatch.setattr(service, "_build_docx_from_markdown", fake_build)

    result = await service.convert_pdf_to_docx("doc_123")

    assert result == {
        "success": True,
        "doc_id": "doc_123",
        "output_path": str(tmp_path / "converted_from_pdf.docx"),
        "mode": "content",
        "figures_embedded": 0,
        "tables_found": 1,
    }
    assert captured["markdown"] == "# Title\n\nHello world"
    assert captured["output_path"] == tmp_path / "converted_from_pdf.docx"


def test_build_docx_from_markdown_reencodes_unrecognized_jpeg(
    monkeypatch, tmp_path: Path
) -> None:
    from docx.document import Document as DocxDocument
    from docx.image.exceptions import UnrecognizedImageError
    from PIL import Image

    image_path = tmp_path / "adobe_header_like.jpeg"
    Image.new("RGB", (16, 16), "white").save(image_path, format="JPEG")
    output_path = tmp_path / "converted.docx"
    service = DocumentService(repository=MagicMock(), pdf_extractor=MagicMock())
    manifest = SimpleNamespace(
        title="",
        assets=SimpleNamespace(
            figures=[SimpleNamespace(path=str(image_path), caption="Figure 1")]
        ),
    )
    original_add_picture = DocxDocument.add_picture
    suffixes: list[str] = []

    def fake_add_picture(self, image_path_or_stream, width=None, height=None):
        suffix = Path(str(image_path_or_stream)).suffix.lower()
        suffixes.append(suffix)
        if suffix in {".jpg", ".jpeg"}:
            raise UnrecognizedImageError
        return original_add_picture(
            self,
            image_path_or_stream,
            width=width,
            height=height,
        )

    monkeypatch.setattr(DocxDocument, "add_picture", fake_add_picture)

    service._build_docx_from_markdown("Body text", manifest, output_path)

    assert output_path.exists()
    assert suffixes == [".jpeg", ".png"]


@pytest.mark.asyncio
async def test_convert_pdf_to_pptx_rejects_fidelity_mode() -> None:
    service = DocumentService(repository=MagicMock(), pdf_extractor=MagicMock())

    result = await service.convert_pdf_to_pptx("doc_123", mode="fidelity")

    assert result["success"] is False
    assert "content mode only" in result["error"]


@pytest.mark.asyncio
async def test_convert_pdf_to_pptx_success(monkeypatch, tmp_path: Path) -> None:
    repository = MagicMock()
    repository.get_doc_dir.return_value = tmp_path
    repository.load_markdown.return_value = "# Slide 1\n\n- bullet"
    repository.load_manifest.return_value = MagicMock(
        title="Deck Title",
        filename="slides.pdf",
        assets=MagicMock(
            figures=[MagicMock(path=tmp_path / "fig1.png", caption="Cap")]
        ),
    )

    service = DocumentService(repository=repository, pdf_extractor=MagicMock())

    captured: dict[str, object] = {}

    def fake_build(markdown, manifest, output_path):
        captured["markdown"] = markdown
        captured["output_path"] = output_path
        return {"total_slides": 3, "figure_slides": 1}

    monkeypatch.setattr(service, "_build_pptx_from_markdown", fake_build)

    result = await service.convert_pdf_to_pptx("doc_123")

    assert result == {
        "success": True,
        "doc_id": "doc_123",
        "output_path": str(tmp_path / "converted_from_pdf.pptx"),
        "mode": "content",
        "slides_created": 3,
        "figure_slides": 1,
    }
    assert captured["markdown"] == "# Slide 1\n\n- bullet"
    assert captured["output_path"] == tmp_path / "converted_from_pdf.pptx"


@pytest.mark.asyncio
async def test_ingest_reports_progress_callback(tmp_path: Path) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 test")

    repository = MagicMock()
    repository.get_doc_dir.return_value = tmp_path / "doc_test"
    repository.save_markdown.return_value = tmp_path / "content.md"
    repository.save_manifest.return_value = None

    extractor = MagicMock()
    extractor.extract_text.return_value = "# Title\n\nHello"
    extractor.get_page_count.return_value = 1
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Paper"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    progress_events: list[tuple[int, int, str, str]] = []

    async def progress_callback(
        step: int, total: int, phase: str, message: str
    ) -> None:
        progress_events.append((step, total, phase, message))

    results = await service.ingest([str(pdf_path)], progress_callback=progress_callback)

    assert len(results) == 1
    assert results[0].success is True
    assert progress_events
    assert progress_events[-1][0] == progress_events[-1][1]
    assert progress_events[-1][2] == "Completed"


@pytest.mark.asyncio
async def test_marker_oom_falls_back_to_pymupdf_with_warning(tmp_path: Path) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 test")

    marker_extractor = MagicMock()
    marker_extractor.parse.side_effect = RuntimeError("CUDA out of memory")

    repository = MagicMock()
    repository.get_doc_dir.return_value = tmp_path / "doc_123"

    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        marker_extractor=marker_extractor,
    )
    service._ingest_single = AsyncMock(
        return_value=IngestResult(
            doc_id="doc_123",
            filename="paper.pdf",
            success=True,
            backend="pymupdf",
        )
    )

    result = await service._ingest_single_with_marker(str(pdf_path))

    assert result.success is True
    assert result.backend == "pymupdf_fallback"
    assert result.warnings
    assert "marker_max_pages_per_chunk=1" in result.warnings[0]
    assert "PyMuPDF fallback" in result.warnings[0]


@pytest.mark.asyncio
async def test_required_marker_oom_does_not_write_fallback_artifacts(
    tmp_path: Path,
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 test")

    marker_extractor = MagicMock()
    marker_extractor.parse.side_effect = RuntimeError("CUDA out of memory")

    repository = MagicMock()
    repository.get_doc_dir.return_value = tmp_path / "doc_123"
    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        marker_extractor=marker_extractor,
    )
    service._ingest_single = AsyncMock(
        return_value=IngestResult(
            doc_id="doc_fallback",
            filename="paper.pdf",
            success=True,
        )
    )

    result = await service._ingest_single_with_marker(
        str(pdf_path),
        require_marker=True,
    )

    assert result.success is False
    assert "marker_max_pages_per_chunk=1" in (result.error or "")
    service._ingest_single.assert_not_awaited()


@pytest.mark.asyncio
async def test_extract_and_save_images_prefers_page_crop_and_spatial_caption(
    tmp_path: Path,
) -> None:
    repository = MagicMock()
    saved_paths: dict[str, Path] = {}

    def fake_save_image(doc_id: str, image_id: str, data: bytes, ext: str) -> Path:
        path = tmp_path / doc_id / "images" / f"{image_id}.{ext}"
        saved_paths[image_id] = path
        return path

    repository.save_image.side_effect = fake_save_image
    raw_left = _test_png_bytes("raw-left")
    raw_right = _test_png_bytes("raw-right")
    crop_left = _test_png_bytes("crop-left-label", size=(288, 272))
    crop_right = _test_png_bytes("crop-right-label", size=(288, 272))

    pdf_extractor = MagicMock()
    pdf_extractor.extract_images.return_value = [
        {
            "page": 4,
            "image_bytes": raw_left,
            "ext": "jpeg",
            "width": 120,
            "height": 90,
            "index_on_page": 1,
            "bbox": [40.0, 80.0, 160.0, 170.0],
            "page_image_bytes": crop_left,
            "page_image_ext": "png",
            "page_crop_bbox": [28.0, 68.0, 172.0, 204.0],
            "page_crop_width": 288,
            "page_crop_height": 272,
            "extraction_strategy": "xobject_page_crop",
        },
        {
            "page": 4,
            "image_bytes": raw_right,
            "ext": "jpeg",
            "width": 120,
            "height": 90,
            "index_on_page": 2,
            "bbox": [260.0, 80.0, 380.0, 170.0],
            "page_image_bytes": crop_right,
            "page_image_ext": "png",
            "page_crop_bbox": [248.0, 68.0, 392.0, 204.0],
            "page_crop_width": 288,
            "page_crop_height": 272,
            "extraction_strategy": "xobject_page_crop",
        },
    ]
    pdf_extractor.extract_figure_captions.return_value = {
        4: [
            {
                "number": "4.1",
                "caption": "Figure 4.1 Left image caption",
                "bbox": [38.0, 178.0, 170.0, 198.0],
            },
            {
                "number": "4.2",
                "caption": "Figure 4.2 Right image caption",
                "bbox": [258.0, 178.0, 390.0, 198.0],
            },
        ]
    }

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        profile=ETLProfile.default(),
    )

    figures = await service._extract_and_save_images(
        "doc_test",
        tmp_path / "source.pdf",
    )

    assert [figure.caption for figure in figures] == [
        "Figure 4.1 Left image caption",
        "Figure 4.2 Right image caption",
    ]
    assert figures[0].path == str(saved_paths["fig_4_1"])
    assert figures[0].raw_path == str(saved_paths["fig_4_1_raw"])
    assert figures[0].ext == "png"
    assert figures[0].width == 288
    assert figures[0].height == 272
    assert figures[0].figure_bbox == [40.0, 80.0, 160.0, 170.0]
    assert figures[0].crop_bbox == [28.0, 68.0, 172.0, 204.0]
    assert figures[0].caption_bbox == [38.0, 178.0, 170.0, 198.0]
    assert figures[0].caption_confidence > 0.8
    assert figures[0].extraction_strategy == "xobject_page_crop"
    assert repository.save_image.call_args_list[0].kwargs == {
        "doc_id": "doc_test",
        "image_id": "fig_4_1_raw",
        "data": raw_left,
        "ext": "jpeg",
    }
    assert repository.save_image.call_args_list[1].kwargs == {
        "doc_id": "doc_test",
        "image_id": "fig_4_1",
        "data": crop_left,
        "ext": "png",
    }


def test_match_caption_by_geometry_prefers_nearest_same_column() -> None:
    figure = {
        "bbox": [260.0, 80.0, 380.0, 170.0],
        "page_crop_bbox": [248.0, 68.0, 392.0, 204.0],
    }
    captions = [
        {
            "caption": "Figure 4.1 Left image caption",
            "bbox": [38.0, 178.0, 170.0, 198.0],
        },
        {
            "caption": "Figure 4.2 Right image caption",
            "bbox": [258.0, 178.0, 390.0, 198.0],
        },
    ]

    match = DocumentService._match_caption_for_image(
        figure,
        captions,
        used_caption_indexes=set(),
    )

    assert match is not None
    assert match["caption"] == "Figure 4.2 Right image caption"
    assert match["bbox"] == [258.0, 178.0, 390.0, 198.0]
    assert match["confidence"] > 0.8


def test_match_caption_by_geometry_rejects_far_caption() -> None:
    figure = {
        "bbox": [40.0, 80.0, 160.0, 170.0],
        "page_crop_bbox": [28.0, 68.0, 172.0, 204.0],
    }
    captions = [
        {
            "caption": "Figure 4.9 Unrelated far caption",
            "bbox": [40.0, 430.0, 170.0, 455.0],
        }
    ]

    match = DocumentService._match_caption_for_image(
        figure,
        captions,
        used_caption_indexes=set(),
    )

    assert match is None


@pytest.mark.asyncio
async def test_extract_and_save_images_does_not_fifo_assign_unmatched_caption(
    tmp_path: Path,
) -> None:
    repository = MagicMock()
    repository.save_image.side_effect = (
        lambda doc_id, image_id, data, ext: tmp_path
        / doc_id
        / "images"
        / f"{image_id}.{ext}"
    )
    crop = _test_png_bytes("crop-label", size=(288, 272))

    pdf_extractor = MagicMock()
    pdf_extractor.extract_images.return_value = [
        {
            "page": 4,
            "image_bytes": _test_png_bytes("raw"),
            "ext": "jpeg",
            "width": 120,
            "height": 90,
            "index_on_page": 1,
            "bbox": [40.0, 80.0, 160.0, 170.0],
            "page_image_bytes": crop,
            "page_image_ext": "png",
            "page_crop_bbox": [28.0, 68.0, 172.0, 204.0],
            "page_crop_width": 288,
            "page_crop_height": 272,
            "extraction_strategy": "xobject_page_crop",
        }
    ]
    pdf_extractor.extract_figure_captions.return_value = {
        4: [
            {
                "number": "4.9",
                "caption": "Figure 4.9 Unrelated far caption",
                "bbox": [40.0, 430.0, 170.0, 455.0],
            }
        ]
    }

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        profile=ETLProfile.default(),
    )

    figures = await service._extract_and_save_images(
        "doc_test", tmp_path / "source.pdf"
    )

    assert len(figures) == 1
    assert figures[0].caption == ""
    assert figures[0].caption_bbox == []
    assert figures[0].caption_confidence == 0.0


@pytest.mark.asyncio
async def test_extract_and_save_images_expands_tiny_xobject_to_caption_anchor_crop(
    monkeypatch, tmp_path: Path
) -> None:
    repository = MagicMock()
    saved_payloads: dict[str, bytes] = {}

    def fake_save_image(doc_id: str, image_id: str, data: bytes, ext: str) -> Path:
        saved_payloads[image_id] = data
        return tmp_path / doc_id / "images" / f"{image_id}.{ext}"

    repository.save_image.side_effect = fake_save_image
    tiny_crop = _test_png_bytes("tiny-probe", size=(80, 180))
    anchored = _test_png_bytes("full-figure-33-1", size=(1291, 953))

    pdf_extractor = MagicMock()
    pdf_extractor.extract_images.return_value = [
        {
            "page": 2,
            "image_bytes": _test_png_bytes("raw-probe"),
            "ext": "jpeg",
            "width": 27,
            "height": 62,
            "index_on_page": 1,
            "bbox": [54.89, 606.564, 82.154, 668.868],
            "page_image_bytes": tiny_crop,
            "page_image_ext": "png",
            "page_crop_bbox": [43.0, 598.564, 558.697, 742.116],
            "page_crop_width": 1032,
            "page_crop_height": 288,
            "extraction_strategy": "xobject_page_crop",
        }
    ]
    pdf_extractor.extract_figure_captions.return_value = {
        2: [
            {
                "number": "33.1",
                "caption": "Figure 33.1 Evolution of transesophageal echocardiography",
                "bbox": [51.0, 707.38, 550.697, 734.116],
            }
        ]
    }

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        profile=ETLProfile.default(),
    )
    monkeypatch.setattr(
        service,
        "_render_page_region_image",
        lambda *_args, **_kwargs: {
            "image": anchored,
            "ext": "png",
            "width": 1291,
            "height": 953,
            "bbox": [43.0, 354.0, 559.0, 738.116],
        },
    )

    figures = await service._extract_and_save_images(
        "doc_test", tmp_path / "source.pdf"
    )

    assert len(figures) == 1
    assert figures[0].extraction_strategy == "caption_anchor_page_crop"
    assert figures[0].width == 1291
    assert figures[0].height == 953
    assert figures[0].crop_bbox == [43.0, 354.0, 559.0, 738.116]
    assert figures[0].figure_bbox == [43.0, 354.0, 559.0, 707.38]
    assert figures[0].caption_bbox == [51.0, 707.38, 550.697, 734.116]
    assert saved_payloads["fig_2_1"] == anchored


@pytest.mark.asyncio
async def test_extract_and_save_images_groups_multiple_xobjects_under_caption(
    monkeypatch,
    tmp_path: Path,
) -> None:
    repository = MagicMock()
    repository.save_image.side_effect = (
        lambda doc_id, image_id, data, ext: tmp_path
        / doc_id
        / "images"
        / f"{image_id}.{ext}"
    )
    component_a = _test_png_bytes("panel-a")
    component_b = _test_png_bytes("panel-b")
    grouped = _test_png_bytes("grouped-caption-crop", size=(500, 360))

    pdf_extractor = MagicMock()
    pdf_extractor.extract_images.return_value = [
        {
            "page": 3,
            "image_bytes": component_a,
            "ext": "jpeg",
            "width": 130,
            "height": 100,
            "index_on_page": 1,
            "bbox": [50.0, 60.0, 180.0, 160.0],
            "page_image_bytes": component_a,
            "page_image_ext": "png",
            "page_crop_bbox": [40.0, 50.0, 190.0, 220.0],
            "page_crop_width": 300,
            "page_crop_height": 340,
            "extraction_strategy": "xobject_page_crop",
        },
        {
            "page": 3,
            "image_bytes": component_b,
            "ext": "jpeg",
            "width": 130,
            "height": 100,
            "index_on_page": 2,
            "bbox": [200.0, 60.0, 330.0, 160.0],
            "page_image_bytes": component_b,
            "page_image_ext": "png",
            "page_crop_bbox": [190.0, 50.0, 340.0, 220.0],
            "page_crop_width": 300,
            "page_crop_height": 340,
            "extraction_strategy": "xobject_page_crop",
        },
        {
            "page": 3,
            "image_bytes": _test_png_bytes("next-figure"),
            "ext": "jpeg",
            "width": 130,
            "height": 100,
            "index_on_page": 3,
            "bbox": [50.0, 230.0, 180.0, 330.0],
            "page_image_bytes": _test_png_bytes("next-figure-crop"),
            "page_image_ext": "png",
            "page_crop_bbox": [40.0, 220.0, 190.0, 390.0],
            "page_crop_width": 300,
            "page_crop_height": 340,
            "extraction_strategy": "xobject_page_crop",
        },
    ]
    pdf_extractor.extract_figure_captions.return_value = {
        3: [
            {
                "number": "42.1",
                "caption": "Figure 42.1 Multipanel caption",
                "bbox": [45.0, 180.0, 360.0, 210.0],
            }
        ]
    }
    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        profile=ETLProfile.default(),
    )
    monkeypatch.setattr(
        service,
        "_render_page_region_image",
        lambda *_args, **_kwargs: {
            "image": grouped,
            "ext": "png",
            "width": 500,
            "height": 360,
            "bbox": [42.0, 52.0, 362.0, 222.0],
        },
    )

    figures = await service._extract_and_save_images(
        "doc_test", tmp_path / "source.pdf"
    )

    assert len(figures) == 2
    assert figures[0].caption == "Figure 42.1 Multipanel caption"
    assert figures[0].extraction_strategy == "caption_group_page_crop"
    assert figures[0].figure_bbox == [50.0, 60.0, 330.0, 160.0]
    assert figures[0].crop_bbox == [42.0, 52.0, 362.0, 222.0]
    assert figures[0].caption_bbox == [45.0, 180.0, 360.0, 210.0]
    assert figures[0].caption_confidence == pytest.approx(1.0)
    assert figures[0].width == 500
    assert figures[0].height == 360
    assert figures[1].figure_bbox == [50.0, 230.0, 180.0, 330.0]
    assert repository.save_image.call_args_list[0].kwargs["data"] == grouped


def test_caption_group_candidates_respect_next_caption_boundary(
    monkeypatch, tmp_path: Path
) -> None:
    service = DocumentService(
        repository=MagicMock(),
        pdf_extractor=MagicMock(),
        profile=ETLProfile.default(),
    )
    monkeypatch.setattr(
        service,
        "_render_page_region_image",
        lambda *_args, **_kwargs: {
            "image": _test_png_bytes("grouped", size=(400, 240)),
            "ext": "png",
            "width": 400,
            "height": 240,
            "bbox": [40.0, 50.0, 340.0, 210.0],
        },
    )

    candidates = [
        {
            "bbox": [50.0, 60.0, 180.0, 160.0],
            "page_crop_bbox": [40.0, 50.0, 190.0, 170.0],
        },
        {
            "bbox": [200.0, 60.0, 330.0, 160.0],
            "page_crop_bbox": [190.0, 50.0, 340.0, 170.0],
        },
        {
            "bbox": [50.0, 245.0, 180.0, 340.0],
            "page_crop_bbox": [40.0, 235.0, 190.0, 350.0],
        },
    ]
    captions = [
        {
            "number": "42.1",
            "caption": "Figure 42.1 Multipanel caption",
            "bbox": [45.0, 180.0, 360.0, 210.0],
        },
        {
            "number": "42.2",
            "caption": "Figure 42.2 Next caption",
            "bbox": [45.0, 230.0, 360.0, 260.0],
        },
    ]

    grouped, consumed, caption_indexes = service._build_caption_group_candidates(
        tmp_path / "source.pdf",
        3,
        candidates,
        captions,
    )

    assert len(grouped) == 1
    assert consumed == {0, 1}
    assert caption_indexes == {0}
    assert grouped[0]["grouped_candidate_count"] == 2


def test_caption_group_candidates_keep_distant_ab_multipanel_together(
    monkeypatch, tmp_path: Path
) -> None:
    service = DocumentService(
        repository=MagicMock(),
        pdf_extractor=MagicMock(),
        profile=ETLProfile.default(),
    )
    monkeypatch.setattr(
        service,
        "_render_page_region_image",
        lambda *_args, **_kwargs: {
            "image": _test_png_bytes("grouped-ab", size=(1042, 640)),
            "ext": "png",
            "width": 1042,
            "height": 640,
            "bbox": [51.0, 40.0, 572.0, 356.0],
        },
    )

    candidates = [
        {
            "bbox": [110.0, 56.0, 376.0, 147.0],
            "page_crop_bbox": [98.0, 38.0, 388.0, 219.0],
        },
        {
            "bbox": [60.0, 144.0, 564.0, 279.0],
            "page_crop_bbox": [52.0, 136.0, 572.0, 356.0],
        },
    ]
    captions = [
        {
            "number": "33.4",
            "caption": "Figure 33.4. (A) Schematic of multiple-beat gated full-volume image acquisition. (B) Creation of the three-dimensional full-volume image from narrow subvolumes.",
            "bbox": [63.0, 312.0, 563.0, 348.0],
        }
    ]

    grouped, consumed, caption_indexes = service._build_caption_group_candidates(
        tmp_path / "source.pdf",
        8,
        candidates,
        captions,
    )

    assert len(grouped) == 1
    assert consumed == {0, 1}
    assert caption_indexes == {0}
    assert grouped[0]["bbox"] == [60.0, 56.0, 564.0, 279.0]
    assert grouped[0]["grouped_candidate_count"] == 2


@pytest.mark.asyncio
async def test_save_marker_images_falls_back_to_bbox_render(
    monkeypatch, tmp_path: Path
) -> None:
    repository = MagicMock()
    repository.save_image.return_value = tmp_path / "fig_7_1.png"
    service = DocumentService(
        repository=repository,
        pdf_extractor=MagicMock(),
        profile=ETLProfile.default(),
    )

    block = MarkerBlock(
        block_id="blk_0001",
        block_type="Figure",
        page=7,
        text="",
        bbox=[10.0, 20.0, 110.0, 180.0],
        metadata={
            "id": "/page/6/Figure/0",
            "caption": "Figure 49.5",
            "source_order": 1,
        },
    )
    parse_result = SimpleNamespace(
        blocks=[block],
        images={},
    )

    def fake_render(*args, **kwargs):
        return (
            b"\x89PNG\r\n\x1a\n"
            b"\x00\x00\x00\rIHDR"
            b"\x00\x00\x00d\x00\x00\x00x\x08\x02\x00\x00\x00"
            b"\x00\x00\x00\x00"
        )

    monkeypatch.setattr(service, "_render_pdf_block_image", fake_render)
    monkeypatch.setattr(service, "_get_image_dimensions", lambda _: (100, 120))

    figures = await service._save_marker_images(
        "doc_test",
        parse_result,
        pdf_path=tmp_path / "source.pdf",
    )

    assert len(figures) == 1
    assert figures[0].id == "fig_7_1"
    assert figures[0].page == 7
    assert figures[0].caption == "Figure 49.5"
    repository.save_image.assert_called_once()


@pytest.mark.asyncio
async def test_ingest_scopes_doc_id_and_page_markers_for_page_ranges(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 original")
    doc_dir = tmp_path / "doc_pages"
    doc_dir.mkdir()

    repository = MagicMock()
    repository.get_doc_dir.return_value = doc_dir
    repository.save_markdown.return_value = tmp_path / "content.md"
    repository.save_manifest.return_value = None

    extractor = MagicMock()

    def fake_get_page_count(path: Path) -> int:
        return 10 if path == pdf_path else 2

    extractor.get_page_count.side_effect = fake_get_page_count
    extractor.extract_text.return_value = (
        "<!-- Page 1 -->\nAlpha\n<!-- Page 2 -->\nBeta"
    )
    extractor.get_toc.return_value = [(1, "Methods", 2)]
    extractor.get_title.return_value = "Paper"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    captured_suffix: dict[str, str] = {}

    def fake_generate(_filename: str, unique_suffix: str):
        captured_suffix["value"] = unique_suffix
        return MagicMock(value="doc_pages")

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate", fake_generate
    )
    monkeypatch.setattr(
        "src.application.document_service.materialize_pdf_page_subset",
        lambda _source_path, output_path, _page_ranges: (
            output_path.write_bytes(b"%PDF-1.4 subset"),
            output_path,
        )[1],
    )

    results = await service.ingest([str(pdf_path)], page_ranges=["3-4"])

    assert results[0].success is True
    assert results[0].pages_processed == 2
    assert captured_suffix["value"].endswith("#pages=3-4")
    repository.save_markdown.assert_called_once_with(
        "doc_pages",
        "<!-- Page 3 -->\nAlpha\n<!-- Page 4 -->\nBeta",
    )
    manifest = repository.save_manifest.call_args.args[0]
    assert manifest.page_count == 10


@pytest.mark.asyncio
async def test_ingest_remaps_table_and_image_pages_for_page_ranges(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 original")
    doc_dir = tmp_path / "doc_pages"
    doc_dir.mkdir()

    repository = MagicMock()
    repository.get_doc_dir.return_value = doc_dir
    repository.save_markdown.return_value = tmp_path / "content.md"
    repository.save_manifest.return_value = None

    extractor = MagicMock()
    extractor.get_page_count.side_effect = lambda path: 10 if path == pdf_path else 2
    extractor.extract_text.return_value = "<!-- Page 1 -->\n# Title"
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Paper"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[MagicMock(page=3)])
    service._extract_tables = AsyncMock(return_value=[MagicMock(page=4)])

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: MagicMock(value="doc_pages"),
    )
    monkeypatch.setattr(
        "src.application.document_service.materialize_pdf_page_subset",
        lambda _source_path, output_path, _page_ranges: (
            output_path.write_bytes(b"%PDF-1.4 subset"),
            output_path,
        )[1],
    )

    await service.ingest([str(pdf_path)], page_ranges=["3-4"])

    service._extract_and_save_images.assert_awaited_once()
    assert service._extract_and_save_images.await_args.kwargs["page_map"] == [3, 4]
    service._extract_tables.assert_awaited_once()
    assert service._extract_tables.await_args.kwargs["page_map"] == [3, 4]


@pytest.mark.asyncio
async def test_ingest_ocr_runs_on_page_range_subset(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 original")
    doc_dir = tmp_path / "doc_pages"
    doc_dir.mkdir()

    repository = MagicMock()
    repository.get_doc_dir.return_value = doc_dir
    repository.save_markdown.return_value = tmp_path / "content.md"
    repository.save_manifest.return_value = None

    extractor = MagicMock()
    extractor.get_page_count.side_effect = lambda path: 10 if path == pdf_path else 2
    extractor.extract_text.return_value = "<!-- Page 1 -->\n# Title"
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Paper"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: MagicMock(value="doc_pages"),
    )
    monkeypatch.setattr(
        "src.application.document_service.materialize_pdf_page_subset",
        lambda _source_path, output_path, _page_ranges: (
            output_path.write_bytes(b"%PDF-1.4 subset"),
            output_path,
        )[1],
    )

    captured: dict[str, Path] = {}

    def fake_ocr(_doc_id: str, source_path: Path, **_kwargs) -> Path:
        captured["source_path"] = source_path
        ocr_path = doc_dir / "ocr_processed.pdf"
        ocr_path.write_bytes(b"%PDF-1.4 ocr")
        return ocr_path

    monkeypatch.setattr(service, "_preprocess_pdf_with_ocr", fake_ocr)

    await service.ingest([str(pdf_path)], ocr_enabled=True, page_ranges=["3-4"])

    assert captured["source_path"] == doc_dir / "selected_pages.pdf"
    extractor.extract_text.assert_called_once_with(doc_dir / "ocr_processed.pdf")


@pytest.mark.asyncio
async def test_ingest_overwrites_stale_original_pdf_copy(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fresh")
    doc_dir = tmp_path / "doc_fixed"
    doc_dir.mkdir()
    (doc_dir / "original.pdf").write_bytes(b"%PDF-1.4 stale")

    repository = MagicMock()
    repository.get_doc_dir.return_value = doc_dir
    repository.save_markdown.return_value = tmp_path / "content.md"
    repository.save_manifest.return_value = None

    extractor = MagicMock()
    extractor.extract_text.return_value = "# Title\n\nHello"
    extractor.get_page_count.return_value = 1
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Paper"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: MagicMock(value="doc_fixed"),
    )

    results = await service.ingest([str(pdf_path)])

    assert results[0].success is True
    assert (doc_dir / "original.pdf").read_bytes() == b"%PDF-1.4 fresh"


@pytest.mark.asyncio
async def test_pymupdf_ingest_persists_searchable_blocks_json(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "chapter.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 original")
    doc_dir = tmp_path / "doc_blocks"
    doc_dir.mkdir()

    repository = MagicMock()
    repository.get_doc_dir.return_value = doc_dir
    repository.save_markdown.return_value = doc_dir / "content.md"
    repository.save_manifest.return_value = None
    repository.save_blocks.side_effect = lambda _doc_id, blocks: (
        (doc_dir / "blocks.json").write_text(
            json.dumps(blocks, ensure_ascii=False), encoding="utf-8"
        ),
        doc_dir / "blocks.json",
    )[1]
    repository.save_citation_index.return_value = doc_dir / "citation_index.jsonl"

    extractor = MagicMock()
    extractor.extract_text.return_value = (
        "<!-- Page 1 -->\n# Airway Management\n\nDifficult mask ventilation predicts "
        "difficult intubation.\n\n<!-- Page 2 -->\n## Rescue Strategy\n\nUse oxygenation first."
    )
    extractor.get_page_count.return_value = 2
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Airway Management"

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: MagicMock(value="doc_blocks"),
    )

    results = await service.ingest([str(pdf_path)])

    assert results[0].success is True
    assert (
        results[0].manifest.source_pdf_sha256
        == sha256(b"%PDF-1.4 original").hexdigest()
    )
    assert results[0].manifest.selected_page_map == []
    blocks_path = doc_dir / "blocks.json"
    assert blocks_path.exists()

    blocks = json.loads(blocks_path.read_text(encoding="utf-8"))
    assert any(block["block_type"] == "SectionHeader" for block in blocks)
    text_blocks = [block for block in blocks if block["block_type"] == "Text"]
    assert text_blocks
    assert any(
        "difficult mask ventilation predicts difficult intubation"
        in block["text"].lower()
        for block in text_blocks
    )
    assert all(
        isinstance((block.get("metadata") or {}).get("line_start"), int)
        and isinstance((block.get("metadata") or {}).get("line_end"), int)
        for block in blocks
    )


@pytest.mark.asyncio
async def test_pymupdf_ingest_persists_pdf_audit_artifacts(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "audit.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 audit")
    repository = FileStorage(tmp_path / "data")

    extractor = MagicMock()
    extractor.extract_text.return_value = "<!-- Page 1 -->\n# Safety\n\nVisible text."
    extractor.get_page_count.return_value = 1
    extractor.get_toc.return_value = []
    extractor.get_title.return_value = "Safety"
    extractor.audit_ai_safety.return_value = {
        "schema_version": "pdf-ai-safety-v1",
        "status": "ok",
        "summary": {"issue_count": 0},
        "issues": [],
    }
    extractor.extract_native_structure.return_value = {
        "schema_version": "pdf-native-structure-v1",
        "backend": "pymupdf",
        "outline": [],
        "pages": [{"page": 1, "width": 300.0, "height": 220.0}],
        "capabilities": {"outline": False},
    }

    service = DocumentService(repository=repository, pdf_extractor=extractor)
    service._extract_and_save_images = AsyncMock(return_value=[])
    service._extract_tables = AsyncMock(return_value=[])

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: SimpleNamespace(value="doc_pdf_audit"),
    )

    results = await service.ingest([str(pdf_path)])
    doc_dir = repository.get_doc_dir("doc_pdf_audit")

    assert results[0].success is True
    assert (doc_dir / "ai_safety_report.json").exists()
    assert (doc_dir / "native_structure.json").exists()
    assert (doc_dir / "segmentation_coverage.json").exists()
    safety = json.loads((doc_dir / "ai_safety_report.json").read_text("utf-8"))
    coverage = json.loads((doc_dir / "segmentation_coverage.json").read_text("utf-8"))
    assert safety["doc_id"] == "doc_pdf_audit"
    assert safety["source_pdf_sha256"] == sha256(b"%PDF-1.4 audit").hexdigest()
    assert safety["analyzed_pdf_sha256"] == sha256(b"%PDF-1.4 audit").hexdigest()
    assert coverage["schema_version"] == "segmentation-coverage-v1"
    assert coverage["metrics"]["segment_count"] > 0
    extractor.audit_ai_safety.assert_called_once_with(doc_dir / "original.pdf")
    extractor.extract_native_structure.assert_called_once_with(doc_dir / "original.pdf")


@pytest.mark.asyncio
async def test_marker_ingest_reports_manifest_counts_and_saves_segmentation(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "paper.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 marker")
    repository = FileStorage(tmp_path / "data")

    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1

    marker_extractor = MagicMock()
    marker_extractor.parse.return_value = MarkerParseResult(
        markdown=(
            "# Abstract\n\n"
            "Marker markdown fallback text.\n\n"
            "| Metric | Value |\n"
            "| --- | --- |\n"
            "| Accuracy | 0.95 |\n"
        ),
        blocks=[
            MarkerBlock(
                block_id="mk_1",
                block_type="MarkdownOutput",
                page=1,
                text="",
                metadata={"source_order": 1},
            )
        ],
        toc=[],
        images={},
        metadata={"title": "Paper"},
        page_count=1,
    )

    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: SimpleNamespace(value="doc_marker"),
    )

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        marker_extractor=marker_extractor,
    )

    result = await service._ingest_single_with_marker(str(pdf_path))
    doc_dir = repository.get_doc_dir("doc_marker")
    status = json.loads(
        (doc_dir / "citation_index.status.json").read_text(encoding="utf-8")
    )
    segmentation = json.loads(
        (doc_dir / "segmentation.json").read_text(encoding="utf-8")
    )
    segments = segmentation["segments"]

    assert result.success is True
    assert result.manifest.source_pdf_sha256 == sha256(b"%PDF-1.4 marker").hexdigest()
    assert result.manifest.selected_page_map == []
    assert result.tables_found == len(result.manifest.assets.tables) == 1
    assert result.sections_found == len(result.manifest.assets.sections) == 1
    assert (doc_dir / "segmentation.json").exists()
    assert (doc_dir / "citation_index.jsonl").exists()
    assert any(segment["segment_type"] == "Table" for segment in segments)
    assert any(segment["segment_type"] == "Section header" for segment in segments)
    assert any(segment.get("asset_id") == "tab_1" for segment in segments)
    assert status["attempted"] is True
    assert status["found"] > 0
    assert status["method"] == "marker_markdown_fallback"
    citation_lines = (
        (doc_dir / "citation_index.jsonl").read_text(encoding="utf-8").splitlines()
    )
    assert citation_lines
    first_span = json.loads(citation_lines[0])
    assert first_span["extraction_backend"] == "marker_markdown_fallback"
    assert any(
        "synthesized markdown-based blocks" in warning for warning in result.warnings
    )


@pytest.mark.asyncio
async def test_marker_ingest_fails_on_empty_markdown(
    monkeypatch, tmp_path: Path
) -> None:
    pdf_path = tmp_path / "empty.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 marker")
    repository = FileStorage(tmp_path / "data")

    pdf_extractor = MagicMock()
    pdf_extractor.get_page_count.return_value = 1
    marker_extractor = MagicMock()
    marker_extractor.parse.return_value = MarkerParseResult(
        markdown="",
        blocks=[],
        toc=[],
        images={},
        metadata={"title": "Empty"},
        page_count=1,
    )
    monkeypatch.setattr(
        "src.application.document_service.DocId.generate",
        lambda *_args: SimpleNamespace(value="doc_empty"),
    )

    service = DocumentService(
        repository=repository,
        pdf_extractor=pdf_extractor,
        marker_extractor=marker_extractor,
    )

    result = await service._ingest_single_with_marker(str(pdf_path))

    assert result.success is False
    assert result.backend == "marker"
    assert "empty markdown" in (result.error or "")
    assert not (repository.get_doc_dir("doc_empty") / "citation_index.jsonl").exists()


def test_markdown_block_builder_does_not_duplicate_table_lines() -> None:
    markdown = "\n".join(
        [
            "# Results",
            "",
            "Intro paragraph.",
            "",
            "| Metric | Value |",
            "| --- | --- |",
            "| Accuracy | 0.95 |",
            "",
            "Closing paragraph.",
        ]
    )
    manifest = DocumentManifest(
        doc_id="doc_123",
        filename="paper.pdf",
        assets=DocumentAssets(
            tables=[
                TableAsset(
                    id="tab_1",
                    page=1,
                    markdown="| Metric | Value |\n| --- | --- |\n| Accuracy | 0.95 |",
                    line_start=4,
                    line_end=7,
                )
            ]
        ),
    )

    blocks = build_markdown_blocks(markdown, manifest)

    table_blocks = [block for block in blocks if block["block_type"] == "Table"]
    text_blocks = [block for block in blocks if block["block_type"] == "Text"]
    assert len(table_blocks) == 1
    assert all("| Metric |" not in block["text"] for block in text_blocks)
