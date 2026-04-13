from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from src.application.dfm_table_bridge import DfmTableBridge
from src.application.docx_service import DocxService
from src.application.table_service import TableService
from src.domain.docx_value_objects import DfmBlockType
from src.infrastructure.excel_renderer import ExcelRenderer
from src.infrastructure.file_storage import FileStorage


def _build_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Drug"
    table.cell(1, 0).text = "Old value"
    doc.save(str(path))


def _build_multi_paragraph_table_docx(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Question"
    table.cell(0, 0).add_paragraph("Second line _____")
    table.cell(0, 1).text = "Answer"
    table.cell(1, 0).text = "Item"
    table.cell(1, 1).text = "Old value"
    doc.save(str(path))


@pytest.fixture
def docx_stack(temp_dir: Path):
    repository = FileStorage(base_dir=temp_dir)
    docx_service = DocxService(repository=repository)
    table_dir = temp_dir / "tables"
    table_service = TableService(
        table_output_dir=table_dir,
        table_renderer=ExcelRenderer(table_dir),
    )
    return repository, docx_service, table_service, DfmTableBridge()


async def _ingest_sample_docx(docx_service: DocxService, temp_dir: Path) -> str:
    sample = temp_dir / "sample.docx"
    _build_sample_docx(sample)
    result = await docx_service.ingest_docx(str(sample))
    assert result["success"] is True
    return str(result["doc_id"])


def _get_first_table_block_id(docx_service: DocxService, doc_id: str) -> str:
    ir = docx_service._load_ir(doc_id)
    assert ir is not None
    for block in ir.blocks:
        if block.block_type == DfmBlockType.TABLE:
            return block.id
    raise AssertionError("Expected at least one table block")


class TestDocxSaveFlowIntegration:
    @pytest.mark.asyncio
    async def test_save_docx_merges_inline_dfm_and_table_context_into_real_docx(
        self, temp_dir: Path, docx_stack
    ) -> None:
        repository, docx_service, table_service, dfm_table_bridge = docx_stack
        doc_id = await _ingest_sample_docx(docx_service, temp_dir)
        block_id = _get_first_table_block_id(docx_service, doc_id)
        output_path = temp_dir / "inline-output.docx"

        with (
            patch("src.presentation.tools.docx_tools.docx_service", docx_service),
            patch("src.presentation.tools.docx_tools.table_service", table_service),
            patch(
                "src.presentation.tools.docx_tools.dfm_table_bridge",
                dfm_table_bridge,
            ),
        ):
            from src.presentation.tools.docx_tools import (
                docx_table_to_context,
                save_docx,
            )

            await docx_table_to_context(doc_id, block_id)
            table_id = next(iter(table_service._tables))
            table_service.update_cell(table_id, 0, "Drug", "Merged inline value")

            dfm_text = await docx_service.get_dfm(doc_id)
            assert dfm_text is not None
            dfm_text = dfm_text.replace("Intro paragraph", "Merged inline paragraph")

            result = await save_docx(doc_id, dfm_text, str(output_path))

        assert "✅" in result
        assert output_path.exists()

        saved = Document(str(output_path))
        assert saved.paragraphs[0].text == "Merged inline paragraph"
        assert saved.tables[0].cell(0, 0).text == "Drug"
        assert saved.tables[0].cell(1, 0).text == "Merged inline value"

        doc_dir = repository.get_doc_dir(doc_id)
        assert "Merged inline value" in (doc_dir / "content.dfm").read_text(
            encoding="utf-8"
        )

    @pytest.mark.asyncio
    async def test_save_docx_merges_split_md_and_table_context_into_real_docx(
        self, temp_dir: Path, docx_stack
    ) -> None:
        repository, docx_service, table_service, dfm_table_bridge = docx_stack
        doc_id = await _ingest_sample_docx(docx_service, temp_dir)
        block_id = _get_first_table_block_id(docx_service, doc_id)
        output_path = temp_dir / "split-output.docx"
        doc_dir = repository.get_doc_dir(doc_id)

        with (
            patch("src.presentation.tools.docx_tools.docx_service", docx_service),
            patch("src.presentation.tools.docx_tools.table_service", table_service),
            patch(
                "src.presentation.tools.docx_tools.dfm_table_bridge",
                dfm_table_bridge,
            ),
        ):
            from src.presentation.tools.docx_tools import (
                docx_table_to_context,
                save_docx,
            )

            await docx_table_to_context(doc_id, block_id)
            table_id = next(iter(table_service._tables))
            table_service.update_cell(table_id, 0, "Drug", "Merged split value")

            md_path = doc_dir / "content.md"
            md_text = md_path.read_text(encoding="utf-8")
            md_path.write_text(
                md_text.replace("Intro paragraph", "Merged split paragraph"),
                encoding="utf-8",
            )

            result = await save_docx(doc_id, output_path=str(output_path), from_md=True)

        assert "✅" in result
        assert output_path.exists()

        saved = Document(str(output_path))
        assert saved.paragraphs[0].text == "Merged split paragraph"
        assert saved.tables[0].cell(1, 0).text == "Merged split value"

        assert "Merged split value" in (doc_dir / "content.dfm").read_text(
            encoding="utf-8"
        )

    @pytest.mark.asyncio
    async def test_save_docx_from_md_preserves_untouched_multi_paragraph_cells(
        self, temp_dir: Path, docx_stack
    ) -> None:
        repository, docx_service, _table_service, _dfm_table_bridge = docx_stack
        sample = temp_dir / "multi-paragraph-table.docx"
        _build_multi_paragraph_table_docx(sample)

        result = await docx_service.ingest_docx(str(sample))
        assert result["success"] is True
        doc_id = str(result["doc_id"])

        doc_dir = repository.get_doc_dir(doc_id)
        md_path = doc_dir / "content.md"
        md_text = md_path.read_text(encoding="utf-8")
        md_path.write_text(
            md_text.replace("Old value", "New value"),
            encoding="utf-8",
        )

        output_path = temp_dir / "multi-paragraph-output.docx"
        save_result = await docx_service.save_docx(
            doc_id,
            output_path=str(output_path),
            from_md=True,
        )

        assert save_result["success"] is True
        saved = Document(str(output_path))

        untouched_cell = saved.tables[0].cell(0, 0)
        assert untouched_cell.paragraphs[0].text == "Question"
        assert untouched_cell.paragraphs[1].text == "Second line _____"
        assert saved.tables[0].cell(1, 1).text == "New value"
