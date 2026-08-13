"""Integration test: mixed PDF/DOCX batch ingestion with real services.

Unlike ``tests/unit/test_mixed_ingest_support.py`` (which mocks
``document_service``/``docx_service`` entirely to test routing logic in
isolation), this exercises the *real* ``DocumentService`` + ``DocxService``
+ ``FileStorage`` stack against real, on-disk PDF and DOCX files, proving the
mixed-batch handler actually produces working documents -- not just a
plausible-looking mocked result.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pymupdf as fitz
import pytest
from docx import Document

from src.application.docx_service import DocxService
from src.domain.etl_profile import ETLProfile
from src.infrastructure.file_storage import FileStorage
from src.infrastructure.pdf_extractor import PyMuPDFExtractor
from src.presentation.tools.mixed_ingest_support import build_mixed_ingest_handler


def _build_real_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page(width=400, height=500)
    page.insert_text((40, 60), "Mixed Batch Integration Test PDF")
    page.insert_text((40, 90), "This paragraph proves real PyMuPDF extraction ran.")
    doc.save(path)
    doc.close()


def _build_real_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Mixed Batch Integration Test DOCX")
    doc.add_paragraph("This paragraph proves real DFM ingestion ran.")
    doc.save(path)


def _assert_document_artifacts_exist(repository, doc: dict) -> None:
    """PDF and DOCX documents have different artifact schemas.

    PDF ingestion produces a `DocumentManifest` (`{doc_id}_manifest.json`);
    DOCX ingestion produces a DFM `content.dfm` + `ir.json` instead -- there
    is no `manifest.json` for DOCX doc_ids. `doc["format"]` (set by the
    mixed-batch handler) is exactly how an agent should decide which
    follow-up tool applies: `inspect_document_manifest`/`document(...)` for
    PDFs, `get_docx_content`/`docx(...)` for DOCX.
    """
    if doc["format"] == "pdf":
        assert repository.load_manifest(doc["doc_id"]) is not None
    elif doc["format"] == "docx":
        dfm_path = repository.get_doc_dir(doc["doc_id"]) / "content.dfm"
        assert dfm_path.exists()
    else:
        pytest.fail(f"unexpected document format: {doc['format']!r}")


@pytest.fixture
def real_stack(tmp_path: Path):
    """A real (non-mocked) DocumentService + DocxService + FileStorage stack."""
    from src.application.document_service import DocumentService

    repository = FileStorage(tmp_path / "data")
    document_service = DocumentService(
        repository=repository,
        pdf_extractor=PyMuPDFExtractor(profile=ETLProfile()),
    )
    docx_service = DocxService(repository=repository)
    return repository, document_service, docx_service


@pytest.mark.asyncio
async def test_mixed_batch_produces_real_loadable_documents(
    real_stack, tmp_path: Path
) -> None:
    """Both a real PDF and a real DOCX must ingest to loadable documents."""
    repository, document_service, docx_service = real_stack

    pdf_path = tmp_path / "study.pdf"
    docx_path = tmp_path / "summary.docx"
    _build_real_pdf(pdf_path)
    _build_real_docx(docx_path)

    reporter = MagicMock()
    reporter.report = AsyncMock()

    handler = build_mixed_ingest_handler(
        [str(pdf_path), str(docx_path)],
        document_service=document_service,
        docx_service=docx_service,
    )
    result = await handler(reporter)

    assert result["success"] is True
    assert result["succeeded"] == 2
    assert result["failed"] == 0
    assert not result["failed_files"]

    docs_by_format = {doc["format"]: doc for doc in result["documents"]}

    pdf_doc_id = docs_by_format["pdf"]["doc_id"]
    assert pdf_doc_id
    pdf_manifest = repository.load_manifest(pdf_doc_id)
    assert pdf_manifest is not None
    assert pdf_manifest.source_engine == "pymupdf"
    pdf_markdown = repository.load_markdown(pdf_doc_id)
    assert "real PyMuPDF extraction ran" in (pdf_markdown or "")

    docx_doc_id = docs_by_format["docx"]["doc_id"]
    assert docx_doc_id
    docx_dfm_path = repository.get_doc_dir(docx_doc_id) / "content.dfm"
    assert docx_dfm_path.exists()
    assert "real DFM ingestion ran" in docx_dfm_path.read_text(encoding="utf-8")

    # Progress must have been reported once per file, in order.
    assert reporter.report.await_count == 2


@pytest.mark.asyncio
async def test_mixed_batch_real_success_survives_a_broken_neighbor(
    real_stack, tmp_path: Path
) -> None:
    """A genuinely corrupt file must not stop the real, valid files from ingesting."""
    repository, document_service, docx_service = real_stack

    pdf_path = tmp_path / "study.pdf"
    _build_real_pdf(pdf_path)
    corrupt_docx_path = tmp_path / "corrupt.docx"
    corrupt_docx_path.write_bytes(b"not a real docx file")
    unsupported_path = tmp_path / "notes.txt"
    unsupported_path.write_text("plain text", encoding="utf-8")

    reporter = MagicMock()
    reporter.report = AsyncMock()

    handler = build_mixed_ingest_handler(
        [str(pdf_path), str(corrupt_docx_path), str(unsupported_path)],
        document_service=document_service,
        docx_service=docx_service,
    )
    result = await handler(reporter)

    assert result["success"] is True
    assert result["succeeded"] == 1
    assert result["failed"] == 2
    assert result["documents"][0]["format"] == "pdf"
    assert repository.load_manifest(result["documents"][0]["doc_id"]) is not None
    failed_files = {item["file"] for item in result["failed_files"]}
    assert str(corrupt_docx_path) in failed_files
    assert str(unsupported_path) in failed_files


@pytest.mark.asyncio
async def test_mixed_batch_all_success_runs_end_to_end_through_a_real_job_service(
    real_stack, tmp_path: Path
) -> None:
    """Ultimate confirmation: real JobService + real handler + real ingest.

    Drives the whole stack an agent actually experiences -- create a
    background job, let it run, then read back progress and the final
    result -- with nothing mocked except job persistence (an in-memory
    store instead of the real JSON-file job store). An all-success batch
    must land as JobStatus.COMPLETED.
    """
    import asyncio

    from src.application.job_service import JobService
    from src.domain.job import Job, JobStatus, JobSummary

    class MemoryJobStore:
        def __init__(self) -> None:
            self.jobs: dict[str, Job] = {}

        async def create(self, job: Job) -> Job:
            self.jobs[job.job_id] = job
            return job

        async def get(self, job_id: str) -> Job | None:
            return self.jobs.get(job_id)

        async def update(self, job: Job) -> Job:
            self.jobs[job.job_id] = job
            return job

        async def list_active(self) -> list[JobSummary]:
            return []

        async def list_all(self, _limit: int = 20) -> list[JobSummary]:
            return []

        async def cleanup_old(self, _max_age_hours: int) -> int:
            return 0

    repository, document_service, docx_service = real_stack
    pdf_path = tmp_path / "study.pdf"
    docx_path = tmp_path / "summary.docx"
    _build_real_pdf(pdf_path)
    _build_real_docx(docx_path)
    file_paths = [str(pdf_path), str(docx_path)]

    store = MemoryJobStore()
    job_service = JobService(job_store=store)
    handler = build_mixed_ingest_handler(
        file_paths,
        document_service=document_service,
        docx_service=docx_service,
    )

    job = await job_service.create_conversion_job(
        operation="ingest_mixed_batch",
        input_files=file_paths,
        handler=handler,
        total_steps=len(file_paths),
    )
    assert job.status == JobStatus.PENDING

    task = job_service._running_tasks[job.job_id]
    await asyncio.wait_for(task, timeout=10)

    finished = await store.get(job.job_id)
    assert finished is not None
    assert finished.status == JobStatus.COMPLETED
    assert finished.progress.percentage == 100.0

    assert finished.result["conversion"]["succeeded"] == 2
    assert finished.result["conversion"]["failed"] == 0
    assert len(finished.result["documents"]) == 2
    assert finished.result["failed_files"] == []
    for doc in finished.result["documents"]:
        _assert_document_artifacts_exist(repository, doc)


@pytest.mark.asyncio
async def test_mixed_batch_partial_failure_marks_job_failed_but_keeps_successes(
    real_stack, tmp_path: Path
) -> None:
    """A batch with any failed file must surface as JobStatus.FAILED.

    This mirrors the existing PDF-only ingest job's convention exactly
    (`_process_ingest_job`: `if failed_files: job.fail(...)`), so agents that
    check `job.status == "completed"` never get a false-positive "all good"
    signal when part of a mixed batch actually failed. The per-file successes
    (and their doc_ids/manifests) must still be fully preserved and usable.
    """
    import asyncio

    from src.application.job_service import JobService
    from src.domain.job import Job, JobStatus, JobSummary

    class MemoryJobStore:
        def __init__(self) -> None:
            self.jobs: dict[str, Job] = {}

        async def create(self, job: Job) -> Job:
            self.jobs[job.job_id] = job
            return job

        async def get(self, job_id: str) -> Job | None:
            return self.jobs.get(job_id)

        async def update(self, job: Job) -> Job:
            self.jobs[job.job_id] = job
            return job

        async def list_active(self) -> list[JobSummary]:
            return []

        async def list_all(self, _limit: int = 20) -> list[JobSummary]:
            return []

        async def cleanup_old(self, _max_age_hours: int) -> int:
            return 0

    repository, document_service, docx_service = real_stack
    pdf_path = tmp_path / "study.pdf"
    docx_path = tmp_path / "summary.docx"
    _build_real_pdf(pdf_path)
    _build_real_docx(docx_path)
    unsupported_path = tmp_path / "notes.txt"
    unsupported_path.write_text("plain text", encoding="utf-8")
    file_paths = [str(pdf_path), str(docx_path), str(unsupported_path)]

    store = MemoryJobStore()
    job_service = JobService(job_store=store)
    handler = build_mixed_ingest_handler(
        file_paths,
        document_service=document_service,
        docx_service=docx_service,
    )

    job = await job_service.create_conversion_job(
        operation="ingest_mixed_batch",
        input_files=file_paths,
        handler=handler,
        total_steps=len(file_paths),
    )

    task = job_service._running_tasks[job.job_id]
    await asyncio.wait_for(task, timeout=10)

    finished = await store.get(job.job_id)
    assert finished is not None
    assert finished.status == JobStatus.FAILED
    assert finished.error and "1" in finished.error

    assert finished.result["conversion"]["succeeded"] == 2
    assert finished.result["conversion"]["failed"] == 1
    assert len(finished.result["documents"]) == 2
    assert len(finished.result["failed_files"]) == 1
    for doc in finished.result["documents"]:
        _assert_document_artifacts_exist(repository, doc)
