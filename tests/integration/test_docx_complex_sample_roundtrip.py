from __future__ import annotations

from pathlib import Path
from typing import Any
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

from src.application.docx_service import DocxService
from src.infrastructure.docx_validator import DocxValidator
from src.infrastructure.file_storage import FileStorage


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _build_complex_docx_fixture(path: Path) -> Path:
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    heading = document.add_heading("Asset-Aware Synthetic DOCX Fixture", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph()
    intro = paragraph.add_run("混合格式段落")
    intro.bold = True
    intro.font.size = Pt(14)
    intro.font.color.rgb = RGBColor(0x16, 0x39, 0x5B)

    emphasis = paragraph.add_run(" with underline")
    emphasis.underline = True
    emphasis.font.name = "Arial"

    suffix = paragraph.add_run("\n第二行說明")
    suffix.font.color.rgb = RGBColor(0x9A, 0x67, 0x1F)

    summary = document.add_paragraph("此測試樣本用於公開 CI，覆蓋巢狀表格與多段文字。")
    summary.paragraph_format.space_after = Pt(6)

    table_one = document.add_table(rows=3, cols=3)
    table_one.style = "Table Grid"
    table_one.cell(0, 0).text = "欄位"
    table_one.cell(0, 1).text = "值"
    table_one.cell(0, 2).text = "狀態"
    _set_cell_shading(table_one.cell(0, 0), "D9EAF4")
    _set_cell_shading(table_one.cell(0, 1), "D9EAF4")
    _set_cell_shading(table_one.cell(0, 2), "D9EAF4")

    merged = table_one.cell(1, 0).merge(table_one.cell(1, 1))
    merged.paragraphs[0].add_run("合併儲存格").bold = True
    table_one.cell(1, 2).text = "保留"

    multi = table_one.cell(2, 0)
    multi.paragraphs[0].text = "多段內容"
    multi.add_paragraph("第二段")
    table_one.cell(2, 1).text = "123"
    table_one.cell(2, 2).text = "完成"

    document.add_paragraph("主表前後仍有普通段落，確認區塊順序不漂移。")

    table_two = document.add_table(rows=2, cols=2)
    table_two.style = "Table Grid"
    table_two.cell(0, 0).text = "區塊"
    table_two.cell(0, 1).text = "內容"
    _set_cell_shading(table_two.cell(0, 0), "F4E3C1")
    _set_cell_shading(table_two.cell(0, 1), "F4E3C1")

    complex_cell = table_two.cell(1, 0)
    complex_cell.paragraphs[0].text = "外層開始"
    complex_cell.add_paragraph("前段備註")

    nested_one = complex_cell.add_table(rows=2, cols=2)
    nested_one.style = "Table Grid"
    nested_one.cell(0, 0).text = "期次"
    nested_one.cell(0, 1).text = "說明"
    nested_one.cell(1, 0).text = "第5次"
    nested_one.cell(1, 1).text = "受試者回診"

    complex_cell.add_paragraph("中段備註")

    nested_two = complex_cell.add_table(rows=2, cols=2)
    nested_two.style = "Table Grid"
    nested_two.cell(0, 0).text = "狀態"
    nested_two.cell(0, 1).text = "結果"
    nested_two.cell(1, 0).text = "核准"
    nested_two.cell(1, 1).text = "待追蹤"

    complex_cell.add_paragraph("外層結束")

    side_cell = table_two.cell(1, 1)
    side_cell.paragraphs[0].text = "右側內容"
    side_cell.add_paragraph("含第二段文字")

    document.add_paragraph("文件尾段，確保 round-trip 前後的非表格內容不受影響。")
    document.save(path)
    return path


@pytest.mark.asyncio
async def test_complex_sample_noop_roundtrip_is_binary_identical(
    temp_dir: Path,
) -> None:
    sample_docx = _build_complex_docx_fixture(temp_dir / "synthetic-complex.docx")
    service = DocxService(repository=FileStorage(base_dir=temp_dir))

    ingest = await service.ingest_docx(str(sample_docx))

    assert ingest["success"] is True
    assert ingest["block_types"]["table"] == 4
    assert ingest["block_types"]["format"] >= 1

    doc_id = str(ingest["doc_id"])
    ir = service._load_ir(doc_id)
    assert ir is not None

    nested_tables = [
        block
        for block in ir.blocks
        if block.block_type.value == "table" and block.parent_cell
    ]
    assert len(nested_tables) == 2
    assert {block.metadata.get("parent_table_id") for block in nested_tables} == {
        "t002"
    }
    assert {block.parent_cell for block in nested_tables} == {"1:0"}

    output = temp_dir / "sample-roundtrip.docx"
    save = await service.save_docx(doc_id, output_path=str(output), from_md=True)

    assert save["success"] is True
    report = DocxValidator().validate(
        temp_dir / doc_id / "original.docx",
        output,
        strict=True,
    )
    assert report.binary_identical is True
    assert report.strict_passed is True


@pytest.mark.asyncio
async def test_complex_sample_nested_table_edit_writes_back(temp_dir: Path) -> None:
    sample_docx = _build_complex_docx_fixture(temp_dir / "synthetic-complex.docx")
    service = DocxService(repository=FileStorage(base_dir=temp_dir))

    ingest = await service.ingest_docx(str(sample_docx))
    assert ingest["success"] is True

    doc_id = str(ingest["doc_id"])
    doc_dir = temp_dir / doc_id
    md_path = doc_dir / "content.md"
    old_value = "第5次"
    new_value = "第5次(修正)"

    md_text = md_path.read_text(encoding="utf-8")
    assert old_value in md_text
    md_path.write_text(md_text.replace(old_value, new_value, 1), encoding="utf-8")

    output = temp_dir / "sample-nested-edit.docx"
    save = await service.save_docx(doc_id, output_path=str(output), from_md=True)

    assert save["success"] is True

    with ZipFile(output) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    text_blob = "|".join(t.text for t in root.findall(".//w:t", ns) if t.text)
    assert new_value in text_blob

    report = DocxValidator().validate(doc_dir / "original.docx", output)
    assert report.structure_score == pytest.approx(1.0)
    assert report.media_score == pytest.approx(1.0)
    assert report.style_score == pytest.approx(1.0)
    assert report.table_score < 1.0
    assert report.table_diffs
