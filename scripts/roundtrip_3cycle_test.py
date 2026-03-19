"""
3-Cycle Round-Trip Fidelity Test

Tests format conversions across multiple cycles to detect cumulative degradation.

Test Scenarios:
1. DOCX → DFM → DOCX (3 cycles) — core round-trip
2. DOCX → ODT → DOCX — cross-format round-trip (OpenDocument Text)
3. DOCX → DOC → DOCX — cross-format round-trip

Each cycle validates 6 dimensions via DocxValidator:
  Structure, Text, Formatting, Tables, Media, Styles

Usage:
    uv run python scripts/roundtrip_3cycle_test.py [path_to_docx]

If no docx is provided, a sample document is generated for testing.
"""

from __future__ import annotations

import asyncio
import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from src.application.docx_service import DocxService
from src.infrastructure.docx_validator import DocxValidator
from src.infrastructure.file_storage import FileStorage


def create_sample_docx(output_path: Path) -> None:
    """Create a sample DOCX with paragraphs, tables, formatting, and styles."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading("Round-Trip Fidelity Test Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Normal paragraphs with various formatting
    p1 = doc.add_paragraph()
    p1.add_run("This is normal text. ")
    run_bold = p1.add_run("This is bold text. ")
    run_bold.bold = True
    run_italic = p1.add_run("This is italic text. ")
    run_italic.italic = True
    run_both = p1.add_run("This is bold+italic.")
    run_both.bold = True
    run_both.italic = True

    # Paragraph with specific font and color
    p2 = doc.add_paragraph()
    run_styled = p2.add_run("Styled text: Arial 14pt Red")
    run_styled.font.name = "Arial"
    run_styled.font.size = Pt(14)
    run_styled.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Heading levels
    doc.add_heading("Section 1: Tables", level=1)
    doc.add_heading("Subsection 1.1: Simple Table", level=2)

    # Simple table
    table1 = doc.add_table(rows=4, cols=3)
    table1.style = "Table Grid"
    headers = ["Name", "Age", "City"]
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
    data = [
        ["Alice", "30", "Taipei"],
        ["Bob", "25", "Tokyo"],
        ["Charlie", "35", "New York"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table1.rows[row_idx].cells[col_idx].text = val

    # Complex table with merged-like content
    doc.add_heading("Subsection 1.2: Data Table", level=2)
    table2 = doc.add_table(rows=3, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["ID", "Parameter", "Value", "Unit"]):
        table2.rows[0].cells[i].text = h
    for i, row in enumerate(
        [["1", "Temperature", "36.5", "°C"], ["2", "Heart Rate", "72", "bpm"]],
        start=1,
    ):
        for j, val in enumerate(row):
            table2.rows[i].cells[j].text = val

    # More sections
    doc.add_heading("Section 2: Lists and Formatting", level=1)

    # Bullet-like paragraphs (using List Bullet style)
    for item in ["First item", "Second item", "Third item"]:
        doc.add_paragraph(item, style="List Bullet")

    # Numbered list
    for item in ["Step one", "Step two", "Step three"]:
        doc.add_paragraph(item, style="List Number")

    # Paragraph with underline and strikethrough
    p3 = doc.add_paragraph()
    run_u = p3.add_run("Underlined text. ")
    run_u.underline = True
    run_s = p3.add_run("Strikethrough text.")
    run_s.font.strike = True

    # Long paragraph to test text wrapping
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
        "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
        "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
        "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
        "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
        "culpa qui officia deserunt mollit anim id est laborum."
    )

    # Chinese text (for CJK round-trip)
    doc.add_heading("Section 3: 中文內容", level=1)
    doc.add_paragraph("這是一段中文測試內容，用來驗證 CJK 字元的往返保真度。")
    doc.add_paragraph("包含特殊符號：①②③ ™ © ® — … 「引號」")

    doc.save(str(output_path))
    print(
        f"  Created sample DOCX: {output_path} ({output_path.stat().st_size:,} bytes)"
    )


async def run_dfm_roundtrip_cycles(
    svc: DocxService,
    validator: DocxValidator,
    original_docx: Path,
    num_cycles: int = 3,
) -> list[dict]:
    """
    Run N cycles of DOCX → DFM → DOCX and validate each cycle.

    Returns a list of per-cycle report dicts.
    """
    results = []
    current_docx = original_docx

    for cycle in range(1, num_cycles + 1):
        print(f"\n{'=' * 60}")
        print(f"  CYCLE {cycle}/{num_cycles}: DOCX → DFM → DOCX")
        print(f"{'=' * 60}")

        # Ingest
        result = await svc.ingest_docx(str(current_docx))
        if not result.get("success"):
            print(f"  ❌ Ingest failed: {result.get('error')}")
            results.append(
                {"cycle": cycle, "success": False, "error": result.get("error")}
            )
            break

        doc_id = result["doc_id"]
        doc_dir = svc.repository.get_doc_dir(doc_id)
        print(f"  doc_id: {doc_id}")
        print(
            f"  blocks: {result.get('total_blocks')}, editable: {result.get('editable_blocks')}"
        )

        # Save without edits (pure round-trip)
        output_docx = doc_dir / f"cycle_{cycle}_output.docx"
        save_result = await svc.save_docx(
            doc_id, output_path=str(output_docx), from_md=True
        )
        if not save_result.get("success"):
            print(f"  ❌ Save failed: {save_result.get('error')}")
            results.append(
                {"cycle": cycle, "success": False, "error": save_result.get("error")}
            )
            break

        # Validate: compare original (cycle 1's original) vs this cycle's output
        report_vs_original = validator.validate(original_docx, output_docx)
        # Also validate: compare previous cycle input vs this cycle's output
        report_vs_prev = validator.validate(current_docx, output_docx)

        cycle_result = {
            "cycle": cycle,
            "success": True,
            "input": str(current_docx),
            "output": str(output_docx),
            "vs_original": report_vs_original.to_dict(),
            "vs_previous": report_vs_prev.to_dict(),
        }
        results.append(cycle_result)

        print(f"\n  vs ORIGINAL: fidelity={report_vs_original.fidelity_score:.1%}")
        print(f"    Structure: {report_vs_original.structure_score:.1%}")
        print(f"    Text:      {report_vs_original.text_score:.1%}")
        print(f"    Format:    {report_vs_original.format_score:.1%}")
        print(f"    Table:     {report_vs_original.table_score:.1%}")
        print(f"    Media:     {report_vs_original.media_score:.1%}")
        print(f"    Style:     {report_vs_original.style_score:.1%}")

        if cycle > 1:
            print(
                f"\n  vs PREVIOUS CYCLE: fidelity={report_vs_prev.fidelity_score:.1%}"
            )

        # Use this cycle's output as next cycle's input
        current_docx = output_docx

    return results


async def run_cross_format_roundtrip(
    svc: DocxService,
    validator: DocxValidator,
    original_docx: Path,
    target_format: str,
) -> dict:
    """
    Run DOCX → {format} → DOCX round-trip and validate.

    target_format: "ods", "doc"
    """
    print(f"\n{'=' * 60}")
    print(f"  CROSS-FORMAT: DOCX → {target_format.upper()} → DOCX")
    print(f"{'=' * 60}")

    # Step 1: Ingest original DOCX
    result = await svc.ingest_docx(str(original_docx))
    if not result.get("success"):
        print(f"  ❌ Ingest failed: {result.get('error')}")
        return {"format": target_format, "success": False, "error": result.get("error")}

    doc_id = result["doc_id"]
    doc_dir = svc.repository.get_doc_dir(doc_id)

    # Step 2: Convert to target format
    convert_method = getattr(svc, f"convert_to_{target_format}")
    target_path = doc_dir / f"cross_format.{target_format}"
    conv_result = await convert_method(doc_id, str(target_path))
    if not conv_result.get("success"):
        print(f"  ❌ DOCX → {target_format.upper()} failed: {conv_result.get('error')}")
        return {
            "format": target_format,
            "success": False,
            "error": conv_result.get("error"),
        }

    print(f"  ✅ DOCX → {target_format.upper()}: {target_path}")

    # Step 3: Convert back to DOCX (ingest the converted file)
    result2 = await svc.ingest_docx(str(target_path))
    if not result2.get("success"):
        print(
            f"  ❌ {target_format.upper()} → DOCX ingest failed: {result2.get('error')}"
        )
        return {
            "format": target_format,
            "success": False,
            "error": result2.get("error"),
        }

    doc_id2 = result2["doc_id"]
    doc_dir2 = svc.repository.get_doc_dir(doc_id2)

    # Step 4: Save back as DOCX
    output_docx = doc_dir2 / f"cross_{target_format}_output.docx"
    save_result = await svc.save_docx(
        doc_id2, output_path=str(output_docx), from_md=True
    )
    if not save_result.get("success"):
        print(f"  ❌ Save failed: {save_result.get('error')}")
        return {
            "format": target_format,
            "success": False,
            "error": save_result.get("error"),
        }

    # Step 5: Validate
    report = validator.validate(original_docx, output_docx)

    print(
        f"\n  DOCX → {target_format.upper()} → DOCX fidelity: {report.fidelity_score:.1%}"
    )
    print(f"    Structure: {report.structure_score:.1%}")
    print(f"    Text:      {report.text_score:.1%}")
    print(f"    Format:    {report.format_score:.1%}")
    print(f"    Table:     {report.table_score:.1%}")
    print(f"    Media:     {report.media_score:.1%}")
    print(f"    Style:     {report.style_score:.1%}")

    if report.text_diffs:
        print(f"\n  Text diffs ({len(report.text_diffs)} found, showing first 5):")
        for d in report.text_diffs[:5]:
            print(f"    [{d.location}]")
            print(f"      orig: {d.original[:80]!r}")
            print(f"      new:  {d.rebuilt[:80]!r}")

    if report.format_diffs:
        print(f"\n  Format diffs ({len(report.format_diffs)} found, showing first 5):")
        for d in report.format_diffs[:5]:
            print(f"    [{d.location}] {d.attribute}: {d.original!r} → {d.rebuilt!r}")

    return {
        "format": target_format,
        "success": True,
        "report": report.to_dict(),
    }


def print_degradation_analysis(results: list[dict]) -> None:
    """Analyze cumulative degradation across cycles."""
    print(f"\n{'=' * 60}")
    print("  DEGRADATION ANALYSIS")
    print(f"{'=' * 60}")

    successful = [r for r in results if r.get("success")]
    if not successful:
        print("  No successful cycles to analyze.")
        return

    print("\n  Fidelity vs. Original across cycles:")
    print("  Cycle | Fidelity | Structure | Text   | Format | Table  | Media  | Style")
    print("  ------|----------|-----------|--------|--------|--------|--------|------")

    for r in successful:
        c = r["cycle"]
        v = r["vs_original"]
        scores = v["scores"]
        print(
            f"  {c:5d} | {v['fidelity_score']:7.1f}% | "
            f"{scores['structure']:8.1f}% | "
            f"{scores['text']:5.1f}% | "
            f"{scores['format']:5.1f}% | "
            f"{scores['table']:5.1f}% | "
            f"{scores['media']:5.1f}% | "
            f"{scores['style']:5.1f}%"
        )

    # Check for degradation
    if len(successful) >= 2:
        first = successful[0]["vs_original"]["fidelity_score"]
        last = successful[-1]["vs_original"]["fidelity_score"]
        delta = last - first

        print(f"\n  Cycle 1 fidelity:     {first:.1f}%")
        print(f"  Cycle {len(successful)} fidelity:     {last:.1f}%")
        print(f"  Cumulative drift:     {delta:+.1f}%")

        if abs(delta) < 0.1:
            print("  ✅ Verdict: NO cumulative degradation — idempotent round-trip")
        elif abs(delta) < 2.0:
            print("  🟡 Verdict: Minor drift detected (< 2%) — acceptable")
        else:
            print("  🔴 Verdict: Significant degradation detected — investigate!")

        # Per-dimension drift
        print("\n  Per-dimension drift (cycle 1 → last):")
        for dim in ["structure", "text", "format", "table", "media", "style"]:
            s1 = successful[0]["vs_original"]["scores"][dim]
            sl = successful[-1]["vs_original"]["scores"][dim]
            d = sl - s1
            indicator = "✅" if abs(d) < 0.1 else ("🟡" if abs(d) < 2.0 else "🔴")
            print(f"    {indicator} {dim:12s}: {s1:6.1f}% → {sl:6.1f}% ({d:+.1f}%)")


async def main() -> None:
    # Determine DOCX path
    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
        if not docx_path.exists():
            print(f"❌ File not found: {docx_path}")
            sys.exit(1)
        use_temp = False
    else:
        # Create a sample DOCX
        tmp_dir = tempfile.mkdtemp(prefix="roundtrip_test_")
        docx_path = Path(tmp_dir) / "sample_test.docx"
        print("No DOCX provided. Creating sample document...")
        create_sample_docx(docx_path)
        use_temp = True

    repo = FileStorage()
    svc = DocxService(repo)
    validator = DocxValidator()

    all_results = {}

    # ── Test 1: 3-cycle DFM round-trip ──
    print("\n" + "═" * 60)
    print("  TEST 1: DOCX → DFM → DOCX (3 cycles)")
    print("═" * 60)
    dfm_results = await run_dfm_roundtrip_cycles(
        svc, validator, docx_path, num_cycles=3
    )
    all_results["dfm_roundtrip"] = dfm_results
    print_degradation_analysis(dfm_results)

    # ── Test 2: Cross-format DOCX → DOC → DOCX ──
    doc_result = await run_cross_format_roundtrip(svc, validator, docx_path, "doc")
    all_results["cross_doc"] = doc_result

    # ── Test 3: Cross-format DOCX → ODT → DOCX ──
    odt_result = await run_cross_format_roundtrip(svc, validator, docx_path, "odt")
    all_results["cross_odt"] = odt_result

    # ── Summary ──
    print(f"\n{'═' * 60}")
    print("  FINAL SUMMARY")
    print(f"{'═' * 60}")

    # DFM round-trip summary
    dfm_ok = all(r.get("success") for r in dfm_results)
    if dfm_ok and dfm_results:
        last_fidelity = dfm_results[-1]["vs_original"]["fidelity_score"]
        print(
            f"  DFM 3-cycle:    {'✅' if last_fidelity >= 95 else '⚠️'} {last_fidelity:.1f}% fidelity after 3 cycles"
        )
    else:
        print("  DFM 3-cycle:    ❌ FAILED")

    # Cross-format summaries
    for name, res in [("DOC", doc_result), ("ODT", odt_result)]:
        if res.get("success"):
            f = res["report"]["fidelity_score"]
            print(f"  DOCX↔{name}:      {'✅' if f >= 80 else '⚠️'} {f:.1f}% fidelity")
        else:
            print(f"  DOCX↔{name}:      ❌ {res.get('error', 'FAILED')}")

    # Save full report
    report_path = Path("data") / "roundtrip_3cycle_report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    with report_path.open("w", encoding="utf-8") as f:
        json.dump(all_results, f, indent=2, ensure_ascii=False, default=str)
    print(f"\n  Full report saved to: {report_path}")

    if use_temp:
        print(f"\n  (Temp dir: {tmp_dir})")


if __name__ == "__main__":
    asyncio.run(main())
