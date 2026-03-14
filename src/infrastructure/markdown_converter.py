"""
Infrastructure Layer - Markdown to DOCX Converter

Converts standalone Markdown text into a .docx file using python-docx.
Supports: headings, paragraphs, bold, italic, strikethrough, code,
inline code, lists (ordered/unordered), tables, horizontal rules,
blockquotes, and images (local file references).
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document as create_document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Type alias ??python-docx's Document() is a factory function;
# the actual class lives in docx.document.Document but is not
# publicly re-exported. We use Any for type annotations.
_Doc = Any

# ============================================================================
# Inline pattern ??order matters (code backtick before bold/italic)
# ============================================================================

_INLINE_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("code", re.compile(r"`([^`]+)`")),
    ("bold_italic", re.compile(r"\*\*\*(.+?)\*\*\*(?!\*)|___(.+?)___(?!_)")),
    ("bold", re.compile(r"\*\*(.+?)\*\*(?!\*)|__(.+?)__(?!_)")),
    ("italic", re.compile(r"\*([^*\n]+)\*(?!\*)|_([^_\n]+)_(?!_)")),
    ("strikethrough", re.compile(r"~~(.+?)~~")),
]


class MarkdownDocxConverter:
    """Convert Markdown text to a .docx Document."""

    def convert(self, md_text: str, output_path: Path) -> Path:
        """
        Parse *md_text* and write the result to *output_path* (.docx).

        Returns the resolved output path.
        """
        doc = create_document()
        self._set_default_style(doc)
        # Normalize line endings (CRLF → LF) before splitting
        lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        i = 0
        while i < len(lines):
            i = self._process_line(doc, lines, i)
        doc.save(str(output_path))
        return output_path

    # ====================================================================
    # Line-level processing
    # ====================================================================

    def _process_line(self, doc: _Doc, lines: list[str], i: int) -> int:
        """Process one logical block starting at line *i*. Return next index."""
        line = lines[i]

        # Blank line ??skip
        if not line.strip():
            return i + 1

        # Fenced code block (```)
        if line.strip().startswith("```"):
            return self._process_code_block(doc, lines, i)

        # Heading (# ??######)
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            return i + 1

        # Horizontal rule
        if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            return i + 1

        # Table (starts with |)
        if line.strip().startswith("|"):
            return self._process_table(doc, lines, i)

        # Blockquote
        if line.strip().startswith(">"):
            return self._process_blockquote(doc, lines, i)

        # Unordered list
        m_ul = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m_ul:
            return self._process_list(doc, lines, i, ordered=False)

        # Ordered list
        m_ol = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m_ol:
            return self._process_list(doc, lines, i, ordered=True)

        # Image (standalone line)
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if m_img:
            self._add_image(doc, m_img.group(2), m_img.group(1))
            return i + 1

        # Normal paragraph
        p = doc.add_paragraph()
        self._add_inline_runs(p, line.strip())
        return i + 1

    # ====================================================================
    # Block-level helpers
    # ====================================================================

    def _process_code_block(self, doc: _Doc, lines: list[str], start: int) -> int:
        """Handle fenced code blocks (``` ... ```)."""
        code_lines: list[str] = []
        i = start + 1
        while i < len(lines):
            if lines[i].strip().startswith("```"):
                i += 1
                break
            code_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return i

    def _process_table(self, doc: _Doc, lines: list[str], start: int) -> int:
        """Parse a markdown pipe table and add a docx table."""
        table_lines: list[str] = []
        i = start
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].strip())
            i += 1

        # Parse rows, skip separator
        rows: list[list[str]] = []
        for tl in table_lines:
            if re.match(r"^\|[\s\-:|]+\|$", tl):
                continue
            cells = [c.strip() for c in tl.split("|")[1:-1]]
            # Restore <br> to newline
            cells = [c.replace("<br>", "\n") for c in cells]
            rows.append(cells)

        if not rows:
            return i

        n_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")

        for ri, row_data in enumerate(rows):
            for ci in range(n_cols):
                cell_text = row_data[ci] if ci < len(row_data) else ""
                cell = table.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                self._add_inline_runs(p, cell_text)

            # Bold header row
            if ri == 0:
                for ci in range(n_cols):
                    for run in table.cell(0, ci).paragraphs[0].runs:
                        run.bold = True

        return i

    def _process_blockquote(self, doc: _Doc, lines: list[str], start: int) -> int:
        """Handle blockquote lines (> ...)."""
        quote_lines: list[str] = []
        i = start
        while i < len(lines) and lines[i].strip().startswith(">"):
            text = re.sub(r"^>\s?", "", lines[i].strip())
            quote_lines.append(text)
            i += 1
        p = (
            doc.add_paragraph(style="Quote")
            if "Quote" in [s.name for s in doc.styles]
            else doc.add_paragraph()
        )
        self._add_inline_runs(p, " ".join(quote_lines))
        p.paragraph_format.left_indent = Inches(0.5)
        return i

    def _process_list(
        self,
        doc: _Doc,
        lines: list[str],
        start: int,
        *,
        ordered: bool,
    ) -> int:
        """Handle ordered/unordered list items."""
        i = start
        item_num = 0
        while i < len(lines):
            line = lines[i]
            if ordered:
                m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
            else:
                m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
            if not m:
                break
            item_num += 1
            indent_level = len(m.group(1)) // 2
            text = m.group(2).strip()
            style = "List Number" if ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
                prefix = f"{item_num}. " if ordered else "• "
                text = prefix + text
            self._add_inline_runs(p, text)
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            i += 1
        return i

    def _add_image(self, doc: _Doc, src: str, alt: str) -> None:
        """Add an image to the document if the file exists."""
        img_path = Path(src)
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(5))
            except Exception:
                logger.warning("Failed to embed image: %s", src)
                doc.add_paragraph(f"[Image: {alt or src}]")
        else:
            doc.add_paragraph(f"[Image: {alt or src}]")

    # ====================================================================
    # Inline formatting
    # ====================================================================

    def _add_inline_runs(self, paragraph: Any, text: str) -> None:
        """Parse inline markdown formatting and add runs to *paragraph*."""
        self._parse_inline(paragraph, text)

    def _parse_inline(self, paragraph: Any, text: str) -> None:
        """Recursively parse inline patterns and emit runs."""
        if not text:
            return

        # Find the earliest match across all patterns
        best_match = None
        best_kind = ""
        best_start = len(text)

        for kind, pattern in _INLINE_PATTERNS:
            m = pattern.search(text)
            if m and m.start() < best_start:
                best_match = m
                best_kind = kind
                best_start = m.start()

        if best_match is None:
            # No inline formatting ??plain text
            if text:
                paragraph.add_run(text)
            return

        # Text before the match
        before = text[: best_match.start()]
        if before:
            paragraph.add_run(before)

        # The matched content
        # Groups may differ per pattern ??pick first non-None group
        content = next((g for g in best_match.groups() if g is not None), "")
        after = text[best_match.end() :]

        if best_kind == "code":
            run = paragraph.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif best_kind == "bold_italic":
            run = paragraph.add_run(content)
            run.bold = True
            run.italic = True
        elif best_kind == "bold":
            run = paragraph.add_run(content)
            run.bold = True
        elif best_kind == "italic":
            run = paragraph.add_run(content)
            run.italic = True
        elif best_kind == "strikethrough":
            run = paragraph.add_run(content)
            run.font.strike = True

        # Continue parsing remainder
        self._parse_inline(paragraph, after)

    # ====================================================================
    # Default styling
    # ====================================================================

    @staticmethod
    def _set_default_style(doc: _Doc) -> None:
        """Set sane default font for the document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
